"""Generate comprehensive project planning documents for AI Visual Dialogue Assistant."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shading)


def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return heading


def add_table_row(table, cells_data, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, "1A56DB")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row


# ============================================================
# Document 1: Project Overview & Architecture
# ============================================================
def generate_overview_doc():
    doc = Document()

    # Title
    title = doc.add_heading('AI 视觉对话助手 — 项目总览与架构设计', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('七牛云 XEngineer 暑期实训营 · 第四批次 · 题目一')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('项目名称：EyeTalk — AI 视觉对话助手\n'
                        '项目定位：基于多模态AI的实时视觉交互应用\n'
                        '技术栈：Web前端 + Python后端 + 七牛云服务 + 多模态大模型')
    run.font.size = Pt(11)

    doc.add_page_break()

    # ---- Section 1: Project Introduction ----
    add_styled_heading(doc, '一、项目简介', 1)

    doc.add_paragraph(
        'EyeTalk 是一款 AI 视觉对话助手应用，用户打开摄像头与麦克风后，AI 能够实时"看到"'
        '摄像头中的视频内容、"听到"用户说的话，并给予自然流畅的语音和文字回应。'
        '项目重点考察多模态AI的综合应用能力，包括视觉理解准确性、语音交互自然度、'
        '以及端云协同的成本控制策略。'
    )

    doc.add_paragraph('')
    add_styled_heading(doc, '1.1 核心亮点', 2)
    highlights = [
        '实时多模态感知：同时处理视频流和音频流，实现"看+听+说"的完整交互闭环',
        '智能帧采样策略：不是每帧都送AI分析，而是基于场景变化检测进行智能采样，大幅降低API调用成本',
        '端云协同架构：轻量级任务（人脸检测、VAD）在端侧完成，重量级任务（视觉理解、对话生成）交给云端AI',
        '自然语音交互：集成ASR+TTS，用户无需打字，全程语音对话',
        '成本可控：多种优化策略确保单次对话成本在可接受范围内',
    ]
    for h in highlights:
        doc.add_paragraph(h, style='List Bullet')

    # ---- Section 2: Architecture ----
    add_styled_heading(doc, '二、系统架构设计', 1)

    add_styled_heading(doc, '2.1 整体架构图（文字描述）', 2)

    arch_text = """
┌─────────────────────────────────────────────────────────────┐
│                    用户端 (浏览器/桌面应用)                      │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐    │
│  │ 摄像头模块 │  │ 麦克风模块 │  │ 视频显示  │  │ 对话界面  │    │
│  └────┬─────┘  └────┬─────┘  └────┬─────┘  └────┬─────┘    │
│       │              │              ▲              ▲          │
│       ▼              ▼              │              │          │
│  ┌──────────────────────────────────────────────────┐       │
│  │              前端控制层 (JavaScript)                │       │
│  │  · 视频帧采集与压缩    · 音频录制与VAD              │       │
│  │  · WebSocket通信       · UI渲染与状态管理            │       │
│  └───────────────────────┬──────────────────────────┘       │
└──────────────────────────┼──────────────────────────────────┘
                           │ WebSocket / HTTP
                           ▼
┌──────────────────────────────────────────────────────────────┐
│                    后端服务层 (Python)                         │
│  ┌───────────┐  ┌───────────┐  ┌───────────┐               │
│  │ WebSocket  │  │ 音频处理   │  │ 视频处理   │               │
│  │ 网关       │  │ 服务       │  │ 服务       │               │
│  └─────┬─────┘  └─────┬─────┘  └─────┬─────┘               │
│        │              │              │                       │
│        ▼              ▼              ▼                       │
│  ┌──────────────────────────────────────────────────┐       │
│  │              AI 调度层                             │       │
│  │  · 多模态请求编排    · 上下文管理                    │       │
│  │  · 成本控制策略      · 响应缓存                     │       │
│  └───────────────────────┬──────────────────────────┘       │
└──────────────────────────┼──────────────────────────────────┘
                           │ API调用
                           ▼
┌──────────────────────────────────────────────────────────────┐
│                    AI 服务层 (云端)                            │
│  ┌───────────┐  ┌───────────┐  ┌───────────┐  ┌──────────┐ │
│  │ 多模态大模型│  │ 语音识别   │  │ 语音合成   │  │ 七牛云    │ │
│  │ (视觉理解) │  │ (ASR)     │  │ (TTS)     │  │ 存储/CDN  │ │
│  └───────────┘  └───────────┘  └───────────┘  └──────────┘ │
└──────────────────────────────────────────────────────────────┘
"""
    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

    # ---- Section 2.2: Module Details ----
    add_styled_heading(doc, '2.2 模块详细说明', 2)

    # Table for modules
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['模块', '职责', '关键技术', '部署位置']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")

    modules = [
        ['摄像头模块', '采集视频流、帧提取、压缩', 'MediaDevices API, Canvas', '前端'],
        ['麦克风模块', '采集音频流、VAD检测', 'Web Audio API, WebRTC', '前端'],
        ['视频帧采样器', '智能帧选择、场景变化检测', 'Canvas像素差异, SSIM', '前端'],
        ['WebSocket网关', '双向实时通信、消息路由', 'FastAPI WebSocket', '后端'],
        ['音频处理服务', 'ASR调用、音频预处理', 'FFmpeg, 七牛云ASR', '后端'],
        ['视频处理服务', '帧解码、格式转换', 'OpenCV, Pillow', '后端'],
        ['AI调度层', '多模态编排、上下文管理', 'Prompt Engineering', '后端'],
        ['多模态大模型', '视觉理解、对话生成', 'GPT-4o / 通义千问VL', '云端'],
        ['ASR服务', '语音转文字', '七牛云ASR / Whisper', '云端'],
        ['TTS服务', '文字转语音', '七牛云TTS / Edge TTS', '云端'],
        ['对象存储', '图片/音频文件存储', '七牛云Kodo', '云端'],
    ]
    for row_data in modules:
        add_table_row(table, row_data)

    # ---- Section 3: Tech Stack ----
    doc.add_page_break()
    add_styled_heading(doc, '三、技术选型', 1)

    add_styled_heading(doc, '3.1 前端技术栈', 2)
    frontend = [
        ('框架', 'React 18 + TypeScript', '组件化开发，类型安全'),
        ('UI库', 'Tailwind CSS + Headless UI', '快速构建响应式界面'),
        ('视频采集', 'MediaDevices API + Canvas', '浏览器原生摄像头访问'),
        ('音频采集', 'Web Audio API + AudioWorklet', '实时音频处理'),
        ('通信', 'WebSocket (原生)', '低延迟双向通信'),
        ('状态管理', 'Zustand', '轻量级状态管理'),
        ('构建工具', 'Vite', '快速开发体验'),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['技术', '选型', '理由']):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")
    for row_data in frontend:
        add_table_row(table, row_data)

    doc.add_paragraph('')
    add_styled_heading(doc, '3.2 后端技术栈', 2)
    backend = [
        ('语言', 'Python 3.11+', '丰富的AI生态库'),
        ('框架', 'FastAPI', '高性能异步框架，原生WebSocket支持'),
        ('音视频处理', 'OpenCV + FFmpeg + Pillow', '帧提取、格式转换'),
        ('ASR集成', '七牛云ASR SDK / OpenAI Whisper', '语音识别'),
        ('TTS集成', '七牛云TTS / Edge-TTS', '语音合成'),
        ('AI模型调用', 'OpenAI SDK / 通义千问SDK', '多模态大模型'),
        ('缓存', 'Redis', '会话上下文、响应缓存'),
        ('部署', 'Docker + Docker Compose', '容器化部署'),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['技术', '选型', '理由']):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")
    for row_data in backend:
        add_table_row(table, row_data)

    doc.add_paragraph('')
    add_styled_heading(doc, '3.3 七牛云服务集成', 2)
    qiniu_services = [
        '七牛云 Kodo 对象存储：存储用户会话中的截图、录音等文件',
        '七牛云 CDN：加速前端静态资源和媒体文件的分发',
        '七牛云 ASR：语音识别服务，将用户语音转为文字',
        '七牛云 TTS：语音合成服务，将AI回复转为语音',
        '七牛云实时音视频（可选）：用于更高质量的音视频传输',
    ]
    for s in qiniu_services:
        doc.add_paragraph(s, style='List Bullet')

    # ---- Section 4: Cost Control ----
    add_styled_heading(doc, '四、成本控制策略（核心竞争力）', 1)

    doc.add_paragraph(
        '成本控制是本项目的核心考察点之一。以下是我们设计的多层次成本控制体系：'
    )

    strategies = [
        ('4.1 视频帧智能采样', [
            '场景变化检测：通过计算相邻帧的像素差异（SSIM/直方图差异），仅在画面有显著变化时才发送帧给AI分析',
            '自适应采样频率：静止场景1-2秒采样一次，动态场景提升到0.5秒一次',
            '人脸触发机制：检测到新面孔或表情变化时优先分析',
            '预期效果：API调用次数降低60-80%',
        ]),
        ('4.2 音频端侧VAD', [
            '前端VAD（Voice Activity Detection）：使用Web Audio API的能量检测 + Silero VAD模型',
            '仅在检测到有效语音时才将音频发送到后端进行ASR',
            '静音段不传输、不处理，节省带宽和算力',
            '预期效果：音频处理成本降低50-70%',
        ]),
        ('4.3 多模态请求合并', [
            '将多帧图像合并到一次请求中（如最近3帧 + 当前帧），减少API调用次数',
            '利用多模态大模型的多图理解能力，一次请求获取更全面的场景描述',
            '对话上下文裁剪：仅保留最近N轮对话，避免token浪费',
        ]),
        ('4.4 响应缓存策略', [
            '相似场景缓存：对相似的视觉场景（余弦相似度>阈值）复用之前的AI分析结果',
            'TTS音频缓存：相同文本的语音合成结果缓存，避免重复调用',
            '会话级缓存：同一会话内的重复查询直接返回缓存结果',
        ]),
        ('4.5 模型分级调用', [
            '第一级（端侧）：人脸检测、场景变化检测 → 零成本',
            '第二级（轻量模型）：简单物体识别、OCR → 低成本模型',
            '第三级（大模型）：复杂场景理解、多轮对话 → 按需调用',
            '根据任务复杂度自动选择合适的模型层级',
        ]),
        ('4.6 流量与带宽优化', [
            '视频帧压缩：JPEG质量80%，分辨率缩放到720p',
            '音频编码：Opus编码，16kHz采样率',
            '使用七牛云CDN加速静态资源，减少源站带宽消耗',
        ]),
    ]

    for title, items in strategies:
        add_styled_heading(doc, title, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    # ---- Section 5: Directory Structure ----
    doc.add_page_break()
    add_styled_heading(doc, '五、项目目录结构', 1)

    dir_structure = """
AI-Visual-Dialogue-Assistant/
├── README.md                      # 项目说明文档
├── docs/                          # 文档目录
│   ├── 01-项目总览与架构设计.md     # 本文档
│   ├── 02-设计文档-用户故事.md      # 用户故事设计文档
│   ├── 03-开发计划与里程碑.md       # 开发计划
│   ├── 04-API接口文档.md           # 接口设计
│   └── 05-成本控制设计文档.md       # 成本控制详细方案
├── src/                           # 源代码
│   ├── frontend/                  # 前端代码
│   │   ├── src/
│   │   │   ├── components/        # React组件
│   │   │   ├── hooks/             # 自定义Hooks
│   │   │   ├── services/          # API服务
│   │   │   ├── stores/            # 状态管理
│   │   │   └── utils/             # 工具函数
│   │   ├── package.json
│   │   └── vite.config.ts
│   ├── backend/                   # 后端代码
│   │   ├── app/
│   │   │   ├── api/               # API路由
│   │   │   ├── services/          # 业务服务
│   │   │   ├── models/            # 数据模型
│   │   │   ├── core/              # 核心配置
│   │   │   └── utils/             # 工具函数
│   │   ├── requirements.txt
│   │   └── main.py
│   └── ai-services/               # AI服务封装
│       ├── vision/                # 视觉理解服务
│       ├── speech/                # 语音服务(ASR/TTS)
│       └── dialogue/              # 对话管理服务
├── config/                        # 配置文件
│   ├── config.yaml                # 主配置
│   └── prompts/                   # Prompt模板
├── tests/                         # 测试代码
│   ├── unit/                      # 单元测试
│   └── integration/               # 集成测试
├── assets/                        # 静态资源
├── docker-compose.yml             # Docker编排
└── Makefile                       # 构建脚本
"""
    p = doc.add_paragraph()
    run = p.add_run(dir_structure)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

    # Save
    path = os.path.join(BASE_DIR, 'docs', '01-项目总览与架构设计.docx')
    doc.save(path)
    print(f"Generated: {path}")


# ============================================================
# Document 2: User Stories & Design Document
# ============================================================
def generate_user_stories_doc():
    doc = Document()

    title = doc.add_heading('AI 视觉对话助手 — 设计文档（用户故事）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run('七牛云 XEngineer 暑期实训营 · 设计文档\n'
                         '包含：计划实现的用户故事 vs 最终实现情况')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ---- User Stories ----
    add_styled_heading(doc, '一、计划实现的用户故事', 1)

    stories = [
        {
            'id': 'US-01',
            'title': '基础视觉对话',
            'priority': 'P0 (必须实现)',
            'description': '作为用户，我希望打开摄像头后，AI能看到我面前的画面并和我对话，以便我向AI展示实物并获得反馈。',
            'acceptance': [
                '用户授权摄像头后，视频画面实时显示在界面上',
                'AI能准确描述摄像头中看到的内容',
                '用户可以通过语音或文字与AI就画面内容进行多轮对话',
                'AI回复延迟 < 3秒',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-02',
            'title': '语音交互',
            'priority': 'P0 (必须实现)',
            'description': '作为用户，我希望通过语音与AI对话，而不需要打字，以便更自然地进行交互。',
            'acceptance': [
                '用户点击麦克风按钮后开始录音',
                '语音识别准确率 > 90%（普通话）',
                'AI回复同时提供文字和语音（TTS）',
                '支持连续对话，无需反复点击',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-03',
            'title': '物体识别与讲解',
            'priority': 'P0 (必须实现)',
            'description': '作为用户，我将某个物体放到摄像头前，AI能识别该物体并提供相关信息。',
            'acceptance': [
                'AI能识别常见物体（书、手机、食物等）',
                '识别后提供物体名称、类别、简要描述',
                '支持中文和英文物体名称',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-04',
            'title': '文字/文档识别（OCR）',
            'priority': 'P1 (应该实现)',
            'description': '作为用户，我将文档或书页放到摄像头前，AI能识别并朗读其中的文字。',
            'acceptance': [
                'AI能识别摄像头中的文字内容',
                '支持中英文混合文字',
                '识别后可语音朗读文字内容',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-05',
            'title': '场景描述与问答',
            'priority': 'P1 (应该实现)',
            'description': '作为用户，我移动摄像头展示不同场景，AI能描述整体场景并回答我的相关问题。',
            'acceptance': [
                'AI能描述场景的整体布局和主要元素',
                '支持"这是什么地方"、"有几个人"等常见问题',
                '场景变化时自动更新描述',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-06',
            'title': '实时翻译',
            'priority': 'P2 (可以实现)',
            'description': '作为用户，我将外文标识或菜单放到摄像头前，AI能识别并翻译成中文。',
            'acceptance': [
                '支持英文、日文、韩文等常见语言',
                '翻译结果以文字+语音形式呈现',
                '保持原文排版结构',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-07',
            'title': '多轮上下文记忆',
            'priority': 'P1 (应该实现)',
            'description': '作为用户，我希望AI能记住之前的对话内容，实现连贯的多轮交互。',
            'acceptance': [
                '至少保持最近10轮对话的上下文',
                '用户可以说"刚才那个"、"再详细说说"等指代性表达',
                '上下文窗口可配置',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-08',
            'title': '智能采样与省流模式',
            'priority': 'P1 (应该实现)',
            'description': '作为用户，我希望在不说话时应用自动降低采样频率，节省流量和费用。',
            'acceptance': [
                '无语音输入时自动降低帧采样频率',
                '检测到语音活动时恢复正常频率',
                '界面上显示当前采样状态和预估费用',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-09',
            'title': '截图与分享',
            'priority': 'P2 (可以实现)',
            'description': '作为用户，我希望能截图保存AI的分析结果，方便后续查看或分享。',
            'acceptance': [
                '一键截图当前画面和AI分析结果',
                '截图保存到七牛云Kodo',
                '生成分享链接',
            ],
            'status': '计划实现',
        },
        {
            'id': 'US-10',
            'title': '多人场景识别',
            'priority': 'P2 (可以实现)',
            'description': '作为用户，当摄像头前有多人时，AI能识别并分别描述每个人的状态。',
            'acceptance': [
                '检测画面中的人数',
                '描述每个人的大致位置和动作',
                '支持简单的表情识别（开心、严肃等）',
            ],
            'status': '计划实现',
        },
    ]

    for story in stories:
        add_styled_heading(doc, f'{story["id"]}：{story["title"]}', 2)

        # Info table
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        info_rows = [
            ('优先级', story['priority']),
            ('用户故事', story['description']),
            ('实现状态', story['status']),
        ]
        for i, (label, value) in enumerate(info_rows):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True if table.rows[i].cells[0].paragraphs[0].runs else False
            table.rows[i].cells[1].text = value

        doc.add_paragraph('')
        p = doc.add_paragraph()
        run = p.add_run('验收标准：')
        run.bold = True
        for ac in story['acceptance']:
            doc.add_paragraph(ac, style='List Bullet')
        doc.add_paragraph('')

    # ---- Implementation Tracking ----
    doc.add_page_break()
    add_styled_heading(doc, '二、实现情况追踪表', 1)

    doc.add_paragraph('（开发完成后在此表中更新实际实现情况）')
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['故事ID', '故事名称', '计划优先级', '实际状态', '备注']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")

    tracking = [
        ['US-01', '基础视觉对话', 'P0', '待开发', ''],
        ['US-02', '语音交互', 'P0', '待开发', ''],
        ['US-03', '物体识别与讲解', 'P0', '待开发', ''],
        ['US-04', '文字/文档识别', 'P1', '待开发', ''],
        ['US-05', '场景描述与问答', 'P1', '待开发', ''],
        ['US-06', '实时翻译', 'P2', '待开发', ''],
        ['US-07', '多轮上下文记忆', 'P1', '待开发', ''],
        ['US-08', '智能采样与省流', 'P1', '待开发', ''],
        ['US-09', '截图与分享', 'P2', '待开发', ''],
        ['US-10', '多人场景识别', 'P2', '待开发', ''],
    ]
    for row_data in tracking:
        add_table_row(table, row_data)

    path = os.path.join(BASE_DIR, 'docs', '02-设计文档-用户故事.docx')
    doc.save(path)
    print(f"Generated: {path}")


# ============================================================
# Document 3: Development Plan & Milestones
# ============================================================
def generate_dev_plan_doc():
    doc = Document()

    title = doc.add_heading('AI 视觉对话助手 — 开发计划与里程碑', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # ---- Phase Overview ----
    add_styled_heading(doc, '一、开发阶段总览', 1)

    doc.add_paragraph(
        '本项目按照"先跑通核心链路，再逐步增强"的策略，分为5个开发阶段。'
        '每个阶段都有明确的交付物和验收标准。'
    )

    phases = [
        {
            'name': '第一阶段：环境搭建与基础链路打通（第1-3天）',
            'tasks': [
                'Day 1：项目初始化',
                '  - 创建前后端项目脚手架（React + FastAPI）',
                '  - 配置开发环境（Python虚拟环境、Node.js、Docker）',
                '  - 申请七牛云开发者账号，获取API Key',
                '  - 配置Git仓库和CI/CD基础',
                '',
                'Day 2：前端摄像头/麦克风模块',
                '  - 实现摄像头权限申请和视频流采集',
                '  - 实现麦克风权限申请和音频流采集',
                '  - 视频预览画面渲染',
                '  - 基础UI布局（视频区 + 对话区 + 控制栏）',
                '',
                'Day 3：后端WebSocket通信',
                '  - FastAPI WebSocket服务搭建',
                '  - 前端WebSocket连接管理',
                '  - 实现视频帧传输（Base64编码）',
                '  - 实现音频流传输',
                '  - 端到端链路测试：前端采集 → 后端接收 → 返回响应',
            ],
            'deliverable': '可运行的前后端应用，能采集音视频并通过WebSocket传输',
        },
        {
            'name': '第二阶段：AI能力集成（第4-7天）',
            'tasks': [
                'Day 4：多模态大模型集成',
                '  - 集成多模态大模型API（GPT-4o / 通义千问VL）',
                '  - 实现图片+文字的多模态请求',
                '  - 编写视觉理解Prompt模板',
                '  - 测试基础的图片描述功能',
                '',
                'Day 5：ASR语音识别集成',
                '  - 集成七牛云ASR / Whisper',
                '  - 实现音频流转文字',
                '  - 处理流式ASR（实时转写）',
                '  - ASR结果与视觉分析结果的融合',
                '',
                'Day 6：TTS语音合成集成',
                '  - 集成TTS服务',
                '  - 实现AI回复的文字转语音',
                '  - 音频流式播放（边生成边播放）',
                '  - 对话流程闭环：看→听→想→说',
                '',
                'Day 7：对话管理',
                '  - 实现多轮对话上下文管理',
                '  - 对话历史存储（Redis）',
                '  - 上下文窗口裁剪策略',
                '  - 完整对话流程联调',
            ],
            'deliverable': '完整的视觉对话功能，能看能听能说',
        },
        {
            'name': '第三阶段：成本优化与智能采样（第8-10天）',
            'tasks': [
                'Day 8：视频帧智能采样',
                '  - 实现场景变化检测（像素差异/SSIM）',
                '  - 自适应采样频率算法',
                '  - 人脸触发采样机制',
                '  - 采样策略可视化（调试用）',
                '',
                'Day 9：音频VAD优化',
                '  - 集成端侧VAD模型',
                '  - 静音检测与过滤',
                '  - 语音活动状态机',
                '  - VAD与帧采样的联动',
                '',
                'Day 10：缓存与模型分级',
                '  - 实现相似场景缓存',
                '  - TTS结果缓存',
                '  - 模型分级调用策略',
                '  - 成本统计与监控面板',
            ],
            'deliverable': '成本优化后的版本，有明确的成本对比数据',
        },
        {
            'name': '第四阶段：功能增强与体验优化（第11-14天）',
            'tasks': [
                'Day 11-12：高级视觉功能',
                '  - OCR文字识别增强',
                '  - 物体识别准确度提升',
                '  - 场景描述能力增强',
                '  - 实时翻译功能',
                '',
                'Day 13：UI/UX优化',
                '  - 对话界面美化',
                '  - 动画效果（录音指示、AI思考动画）',
                '  - 响应式布局适配',
                '  - 错误处理与用户提示',
                '',
                'Day 14：截图与分享',
                '  - 截图功能实现',
                '  - 七牛云Kodo上传',
                '  - 分享链接生成',
            ],
            'deliverable': '功能完善、体验良好的版本',
        },
        {
            'name': '第五阶段：测试、文档与交付（第15-17天）',
            'tasks': [
                'Day 15：测试与Bug修复',
                '  - 单元测试编写',
                '  - 集成测试',
                '  - 性能测试（延迟、并发）',
                '  - Bug修复',
                '',
                'Day 16：文档完善',
                '  - 设计文档更新（实际实现情况）',
                '  - API接口文档',
                '  - 部署文档',
                '  - 用户使用手册',
                '',
                'Day 17：部署与演示',
                '  - Docker镜像构建',
                '  - 部署到服务器',
                '  - 演示视频录制',
                '  - 项目提交',
            ],
            'deliverable': '可部署运行的完整项目 + 完整文档 + 演示视频',
        },
    ]

    for phase in phases:
        add_styled_heading(doc, phase['name'], 2)
        for task in phase['tasks']:
            if task == '':
                doc.add_paragraph('')
            elif task.startswith('  '):
                doc.add_paragraph(task.strip(), style='List Bullet')
            else:
                p = doc.add_paragraph()
                run = p.add_run(task)
                run.bold = True
        p = doc.add_paragraph()
        run = p.add_run(f'阶段交付物：{phase["deliverable"]}')
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        doc.add_paragraph('')

    # ---- Milestone Table ----
    doc.add_page_break()
    add_styled_heading(doc, '二、里程碑节点', 1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['里程碑', '时间节点', '关键交付', '验收标准']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")

    milestones = [
        ['M1', '第3天', '基础链路打通', '音视频能从浏览器传到后端'],
        ['M2', '第7天', '完整对话功能', '能看到、能听到、能回答'],
        ['M3', '第10天', '成本优化版本', 'API调用次数降低50%以上'],
        ['M4', '第14天', '功能增强版本', '全部P0+P1用户故事完成'],
        ['M5', '第17天', '最终交付', '完整项目+文档+演示'],
    ]
    for row_data in milestones:
        add_table_row(table, row_data)

    # ---- Risk Management ----
    add_styled_heading(doc, '三、风险与应对', 1)

    risks = [
        ('多模态大模型API延迟高', '使用流式响应；前端显示"思考中"动画；设置超时重试机制'),
        ('浏览器摄像头兼容性', '优先支持Chrome；提供降级方案（仅语音/仅文字）'),
        ('ASR识别准确率不足', '支持文字输入作为备选；提供编辑功能修正识别结果'),
        ('成本超出预期', '严格实施帧采样和VAD策略；设置每日费用上限告警'),
        ('七牛云服务不可用', '准备备选方案（Whisper本地ASR、Edge-TTS等）'),
    ]

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    for i, h in enumerate(['风险', '应对策略']):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")
    for risk, strategy in risks:
        add_table_row(table, [risk, strategy])

    path = os.path.join(BASE_DIR, 'docs', '03-开发计划与里程碑.docx')
    doc.save(path)
    print(f"Generated: {path}")


# ============================================================
# Document 4: Cost Control Design Document
# ============================================================
def generate_cost_control_doc():
    doc = Document()

    title = doc.add_heading('AI 视觉对话助手 — 成本控制设计文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run('七牛云 XEngineer 暑期实训营 · 设计文档\n'
                         '内容：成本控制技巧的设计思路与实际采用情况')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    add_styled_heading(doc, '一、成本分析', 1)

    doc.add_paragraph(
        '在AI视觉对话场景中，主要的运营成本来自以下几个方面：'
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['成本项', '计费单位', '预估单价', '占比']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")

    costs = [
        ['多模态大模型API', '按Token/图片张数', '约0.01-0.1元/张', '50-60%'],
        ['ASR语音识别', '按时长', '约0.006元/秒', '15-20%'],
        ['TTS语音合成', '按字符数', '约0.002元/百字', '5-10%'],
        ['带宽流量', '按流量', '约0.5元/GB', '10-15%'],
        ['对象存储', '按容量+请求', '极低', '2-5%'],
    ]
    for row_data in costs:
        add_table_row(table, row_data)

    doc.add_paragraph('')
    doc.add_paragraph(
        '可以看到，多模态大模型API调用是最大的成本来源，因此成本控制的重点在于减少不必要的API调用。'
    )

    # ---- Strategies ----
    add_styled_heading(doc, '二、成本控制技巧设计', 1)

    add_styled_heading(doc, '2.1 想到的控制技巧（设计阶段）', 2)

    all_strategies = [
        {
            'name': '视频帧智能采样',
            'type': '核心策略',
            'description': '不每帧都发送给AI分析，而是智能选择有价值的帧',
            'approaches': [
                '场景变化检测：计算相邻帧的像素差异（MSE/SSIM），仅差异超过阈值时才采样',
                '定时+触发双模式：定时采样（如每2秒）+ 事件触发（检测到人脸/运动时）',
                '自适应频率：根据场景动态变化速度自动调整采样间隔',
            ],
            'planned': True,
            'implemented': True,
            'effect': '预计降低60-80%的视觉API调用',
        },
        {
            'name': '端侧VAD（语音活动检测）',
            'type': '核心策略',
            'description': '在前端检测是否有有效语音，静音段不传输',
            'approaches': [
                '能量阈值检测：音频能量低于阈值时判定为静音',
                'Silero VAD模型：使用轻量级VAD模型在浏览器端运行',
                '语音前后缓冲：保留语音前后各200ms的缓冲区',
            ],
            'planned': True,
            'implemented': True,
            'effect': '预计降低50-70%的ASR调用',
        },
        {
            'name': '多模态请求合并',
            'type': '优化策略',
            'description': '将多个信息合并到一次请求中',
            'approaches': [
                '多图合并：将最近N帧图片合并到一次请求中',
                '文字+图片合并：将用户的语音文字和当前帧一起发送',
                '上下文压缩：用摘要替代完整的对话历史',
            ],
            'planned': True,
            'implemented': True,
            'effect': '预计降低30-40%的API调用次数',
        },
        {
            'name': '响应缓存',
            'type': '优化策略',
            'description': '对相似请求复用之前的响应结果',
            'approaches': [
                '图片相似度缓存：对相似场景（SSIM>0.95）复用分析结果',
                'TTS文本缓存：相同文本的语音合成结果缓存24小时',
                '对话缓存：完全相同的查询直接返回缓存',
            ],
            'planned': True,
            'implemented': True,
            'effect': '预计降低10-20%的重复调用',
        },
        {
            'name': '模型分级调用',
            'type': '优化策略',
            'description': '根据任务复杂度选择不同级别的模型',
            'approaches': [
                '简单任务用小模型：如物体分类用轻量模型，复杂理解用大模型',
                '两阶段调用：先用小模型判断是否需要大模型介入',
                '本地模型优先：能用本地模型解决的不调云端',
            ],
            'planned': True,
            'implemented': True,
            'effect': '预计降低20-30%的大模型调用',
        },
        {
            'name': '传输优化',
            'type': '基础策略',
            'description': '减少传输的数据量',
            'approaches': [
                '图片压缩：JPEG质量80%，分辨率不超过720p',
                '音频编码：使用Opus编码，16kHz采样率',
                '增量传输：仅传输变化的帧区域',
            ],
            'planned': True,
            'implemented': True,
            'effect': '预计降低40-60%的带宽消耗',
        },
        {
            'name': '对话上下文裁剪',
            'type': '基础策略',
            'description': '控制对话历史的长度，减少token消耗',
            'approaches': [
                '滑动窗口：仅保留最近N轮对话',
                '摘要压缩：对早期对话生成摘要替代原文',
                '重要性保留：标记重要的对话轮次，优先保留',
            ],
            'planned': True,
            'implemented': True,
            'effect': '预计降低20-30%的token消耗',
        },
        {
            'name': '费用监控与限流',
            'type': '保障策略',
            'description': '设置费用上限，防止意外超支',
            'approaches': [
                '实时费用统计：每次API调用记录费用',
                '日/月费用上限告警',
                '用户级配额控制',
                '降级策略：费用接近上限时自动切换到低成本模式',
            ],
            'planned': True,
            'implemented': True,
            'effect': '保障成本可控',
        },
    ]

    for strategy in all_strategies:
        add_styled_heading(doc, f'{strategy["name"]}（{strategy["type"]}）', 3)
        doc.add_paragraph(strategy['description'])
        for approach in strategy['approaches']:
            doc.add_paragraph(approach, style='List Bullet')
        p = doc.add_paragraph()
        run = p.add_run(f'预期效果：{strategy["effect"]}')
        run.italic = True
        doc.add_paragraph('')

    # ---- Implementation Status ----
    doc.add_page_break()
    add_styled_heading(doc, '三、实际采用情况', 1)

    doc.add_paragraph(
        '以下表格记录了各成本控制技巧的实际实施状态：'
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    headers = ['策略', '计划', '实际采用', '实际效果', '备注']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")

    status_rows = [
        ['视频帧智能采样', '✓', '待填写', '待填写', ''],
        ['端侧VAD', '✓', '待填写', '待填写', ''],
        ['多模态请求合并', '✓', '待填写', '待填写', ''],
        ['响应缓存', '✓', '待填写', '待填写', ''],
        ['模型分级调用', '✓', '待填写', '待填写', ''],
        ['传输优化', '✓', '待填写', '待填写', ''],
        ['对话上下文裁剪', '✓', '待填写', '待填写', ''],
        ['费用监控与限流', '✓', '待填写', '待填写', ''],
    ]
    for row_data in status_rows:
        add_table_row(table, row_data)

    doc.add_paragraph('')
    doc.add_paragraph('注：开发完成后，请在此表中更新实际采用情况和效果数据。')

    # ---- Cost Comparison ----
    add_styled_heading(doc, '四、成本对比（优化前 vs 优化后）', 1)

    doc.add_paragraph('（开发完成后填入实际数据）')
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['指标', '优化前（预估）', '优化后（实测）', '降幅']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")

    comparison = [
        ['每分钟API调用次数', '60次/分钟', '待测量', ''],
        ['每分钟Token消耗', '约10000 tokens', '待测量', ''],
        ['每分钟带宽消耗', '约10MB', '待测量', ''],
        ['单次对话成本（5分钟）', '约1-2元', '待测量', ''],
    ]
    for row_data in comparison:
        add_table_row(table, row_data)

    path = os.path.join(BASE_DIR, 'docs', '05-成本控制设计文档.docx')
    doc.save(path)
    print(f"Generated: {path}")


# ============================================================
# Document 5: API Interface Document
# ============================================================
def generate_api_doc():
    doc = Document()

    title = doc.add_heading('AI 视觉对话助手 — API接口文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    add_styled_heading(doc, '一、WebSocket接口', 1)

    add_styled_heading(doc, '1.1 连接地址', 2)
    doc.add_paragraph('ws://localhost:8000/ws/dialogue')

    add_styled_heading(doc, '1.2 消息格式', 2)
    doc.add_paragraph('所有消息使用JSON格式传输：')

    msg_types = [
        ('客户端 → 服务端', [
            ('video_frame', '{"type": "video_frame", "data": "<base64_image>", "timestamp": 1234567890}'),
            ('audio_chunk', '{"type": "audio_chunk", "data": "<base64_audio>", "timestamp": 1234567890}'),
            ('text_message', '{"type": "text", "content": "这是什么物体？"}'),
            ('control', '{"type": "control", "action": "start|stop|pause"}'),
        ]),
        ('服务端 → 客户端', [
            ('ai_response', '{"type": "ai_response", "text": "这是一个苹果", "audio": "<base64_tts>", "visual_tags": ["apple", "fruit"]}'),
            ('status', '{"type": "status", "state": "listening|thinking|speaking", "sampling_rate": 1}'),
            ('error', '{"type": "error", "code": 500, "message": "服务暂不可用"}'),
        ]),
    ]

    for direction, messages in msg_types:
        p = doc.add_paragraph()
        run = p.add_run(direction)
        run.bold = True
        for msg_type, example in messages:
            doc.add_paragraph(f'消息类型：{msg_type}', style='List Bullet')
            p = doc.add_paragraph()
            run = p.add_run(f'  {example}')
            run.font.size = Pt(9)
            run.font.name = 'Consolas'

    add_styled_heading(doc, '二、HTTP接口', 1)

    apis = [
        ('POST /api/session/create', '创建新的对话会话', '{"user_id": "optional"}', '{"session_id": "uuid", "status": "ok"}'),
        ('GET /api/session/{id}/history', '获取会话对话历史', '-', '{"messages": [...]}'),
        ('POST /api/session/{id}/screenshot', '上传截图到七牛云', '{"image": "base64"}', '{"url": "https://...", "key": "..."}'),
        ('GET /api/cost/stats', '获取费用统计', '-', '{"today_cost": 1.23, "month_cost": 45.67}'),
        ('GET /api/health', '健康检查', '-', '{"status": "ok", "version": "1.0.0"}'),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    headers = ['接口', '说明', '请求参数', '响应']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")
    for api in apis:
        add_table_row(table, api)

    path = os.path.join(BASE_DIR, 'docs', '04-API接口文档.docx')
    doc.save(path)
    print(f"Generated: {path}")


# ============================================================
# Main
# ============================================================
if __name__ == '__main__':
    generate_overview_doc()
    generate_user_stories_doc()
    generate_dev_plan_doc()
    generate_cost_control_doc()
    generate_api_doc()
    print("\nAll documents generated successfully!")
