"""Generate 1.5-day sprint plan with Claude Code prompts for each step."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color):
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading_elm.append(shading)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return h


def make_header_row(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")


def add_row(table, cells_data):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(text))
        run.font.size = Pt(10)
    return row


def add_code_block(doc, code):
    """Add a styled code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Light gray background via shading
    shading = run._element.get_or_add_rPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F5F5F5', qn('w:val'): 'clear'
    })
    shading.append(shading_elm)
    return p


def add_prompt_block(doc, prompt_text):
    """Add a Claude prompt block with blue left border style."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Icon prefix
    run = p.add_run("🤖 Claude Prompt:\n")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x21, 0xA8)
    # The actual prompt
    run2 = p.add_run(prompt_text)
    run2.font.size = Pt(10)
    run2.font.name = 'Consolas'
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def generate():
    doc = Document()

    # ===== TITLE =====
    title = doc.add_heading('EyeTalk — 一天半极速冲刺开发计划', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        '七牛云 XEngineer 暑期实训营 · 第四批次 · 题目一：AI 视觉对话助手\n'
        'AI 后端：DeepSeek API | 开发辅助：Claude Code\n'
        '总时间：1.5天（约12-14小时有效开发时间）'
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ====================================================================
    # SECTION 0: 总体策略
    # ====================================================================
    add_heading_styled(doc, '〇、总体策略与技术选型', 1)

    doc.add_paragraph(
        '一天半时间非常紧张，必须采用"最小可用 → 逐步增强"的策略。'
        '核心原则：先跑通主链路，再加优化，绝不一开始就追求完美。'
    )

    add_heading_styled(doc, '0.1 技术选型（精简版）', 2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    make_header_row(table, ['层级', '选型', '理由'])
    rows = [
        ['前端', '原生 HTML + JavaScript + CSS', '零构建工具，打开即用，最快上手'],
        ['后端', 'Python FastAPI', '原生异步+WebSocket，AI生态最好'],
        ['AI对话', 'DeepSeek API（deepseek-chat）', '你指定的API，兼容OpenAI格式'],
        ['视觉理解', 'DeepSeek Vision API', 'DeepSeek-VL / 多模态接口'],
        ['语音识别(ASR)', '浏览器 Web Speech API', '免费、零延迟、无需后端'],
        ['语音合成(TTS)', '浏览器 SpeechSynthesis API', '免费、无需后端调用'],
        ['实时通信', 'WebSocket', '低延迟双向通信'],
        ['部署', '本地运行即可', '实训营演示用，无需云部署'],
    ]
    for r in rows:
        add_row(table, r)

    doc.add_paragraph('')
    add_heading_styled(doc, '0.2 极简架构', 2)

    arch = """
用户浏览器
  ├─ 摄像头 (getUserMedia) → Canvas截图 → Base64
  ├─ 麦克风 (Web Speech API) → 文字
  └─ WebSocket ──────────────→ 后端 FastAPI
                                    ├─ 接收图片+文字
                                    ├─ 调用 DeepSeek Vision API
                                    └─ 返回AI回复 → 前端显示+TTS朗读
"""
    p = doc.add_paragraph()
    run = p.add_run(arch)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

    add_heading_styled(doc, '0.3 时间分配', 2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    make_header_row(table, ['时间段', '目标', '时长'])
    time_rows = [
        ['Day 1 上午', '环境搭建 + 前端摄像头/麦克风 + UI骨架', '3-4小时'],
        ['Day 1 下午', '后端WebSocket + DeepSeek API接入 + 基础对话', '3-4小时'],
        ['Day 1 晚上', '视觉理解 + 语音交互闭环 + 联调', '2-3小时'],
        ['Day 2 上午', '成本优化(帧采样/VAD) + 高级功能 + 美化', '3-4小时'],
    ]
    for r in time_rows:
        add_row(table, r)

    doc.add_page_break()

    # ====================================================================
    # STEP 1: Environment Setup
    # ====================================================================
    add_heading_styled(doc, '第一步：环境搭建与项目初始化', 1)

    p = doc.add_paragraph()
    run = p.add_run('⏱ 预计耗时：30分钟')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE0, 0x4B, 0x00)

    add_heading_styled(doc, '1.1 要做的事情', 2)
    tasks = [
        '创建项目文件夹结构',
        '创建后端 Python 虚拟环境，安装依赖',
        '创建前端 HTML 骨架文件',
        '申请 DeepSeek API Key（https://platform.deepseek.com/）',
        '配置 .env 文件存放 API Key',
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    add_heading_styled(doc, '1.2 目录结构', 2)
    structure = """
eye-talk/
├── backend/
│   ├── main.py              # FastAPI 主入口
│   ├── ai_service.py        # DeepSeek API 调用封装
│   ├── requirements.txt     # Python 依赖
│   └── .env                 # API Key 配置
├── frontend/
│   ├── index.html           # 主页面
│   ├── app.js               # 核心逻辑
│   ├── style.css            # 样式
│   └── vad.js               # 端侧语音检测(后加)
└── README.md
"""
    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

    add_heading_styled(doc, '1.3 问 Claude 的语句', 2)

    add_prompt_block(doc,
        '帮我创建一个 AI 视觉对话助手项目。技术栈：前端用原生 HTML+JS+CSS，'
        '后端用 Python FastAPI。请帮我：\n'
        '1. 创建项目目录结构（eye-talk/backend/ 和 eye-talk/frontend/）\n'
        '2. 在 backend/ 下创建 requirements.txt，包含 fastapi, uvicorn, '
        'websockets, python-dotenv, openai, pillow, httpx\n'
        '3. 在 backend/ 下创建 .env 文件模板，包含 DEEPSEEK_API_KEY=your_key_here\n'
        '4. 在 frontend/ 下创建 index.html 基础骨架，包含视频预览区、对话消息区、控制按钮\n'
        '5. 创建 README.md 说明如何启动项目'
    )

    add_prompt_block(doc,
        '在 backend/ 目录下创建一个最简单的 FastAPI main.py，要求：\n'
        '- 包含一个 GET / 健康检查接口\n'
        '- 包含一个 WebSocket /ws 接口，接收消息后原样返回(echo)\n'
        '- 从 .env 读取 DEEPSEEK_API_KEY\n'
        '- 使用 uvicorn 启动，端口 8000\n'
        '- 同时提供启动命令：pip install -r requirements.txt && python main.py'
    )

    doc.add_page_break()

    # ====================================================================
    # STEP 2: Frontend Camera & Mic
    # ====================================================================
    add_heading_styled(doc, '第二步：前端摄像头与麦克风模块', 1)

    p = doc.add_paragraph()
    run = p.add_run('⏱ 预计耗时：1-1.5小时')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE0, 0x4B, 0x00)

    add_heading_styled(doc, '2.1 要做的事情', 2)
    tasks = [
        '实现摄像头权限申请与视频流展示',
        '实现 Canvas 截帧功能（将视频帧转为 Base64 图片）',
        '实现麦克风权限申请',
        '使用 Web Speech API 做语音识别（浏览器端ASR，免费）',
        '基础 UI：视频画面 + 录音按钮 + 消息列表',
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    add_heading_styled(doc, '2.2 关键功能说明', 2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    make_header_row(table, ['功能', '实现方式', '关键API'])
    func_rows = [
        ['摄像头采集', 'navigator.mediaDevices.getUserMedia({video:true})', 'MediaDevices API'],
        ['视频预览', '<video> 元素绑定 stream', 'HTMLVideoElement'],
        ['截帧', 'Canvas drawImage + toDataURL("image/jpeg", 0.8)', 'Canvas API'],
        ['麦克风采集', 'getUserMedia({audio:true})', 'MediaDevices API'],
        ['语音识别', 'webkitSpeechRecognition / SpeechRecognition', 'Web Speech API'],
        ['语音合成', 'speechSynthesis.speak(utterance)', 'SpeechSynthesis API'],
    ]
    for r in func_rows:
        add_row(table, r)

    add_heading_styled(doc, '2.3 问 Claude 的语句', 2)

    add_prompt_block(doc,
        '现在帮我编写 frontend/index.html 和 frontend/app.js，实现以下功能：\n\n'
        '1. 页面布局：\n'
        '   - 左侧：视频预览区（显示摄像头画面），下方放"开始摄像头"和"停止"按钮\n'
        '   - 右侧：对话消息区（聊天泡泡样式），底部放"按住说话"按钮和文字输入框\n'
        '   - 顶部：标题栏"EyeTalk - AI视觉对话助手"\n'
        '   - 样式用内联 CSS 或 style.css，深色主题，圆角卡片风格\n\n'
        '2. 摄像头功能：\n'
        '   - 点击"开始摄像头"按钮，申请摄像头权限并显示预览\n'
        '   - 提供 captureFrame() 函数，将当前视频帧截取为 Base64 JPEG（质量0.8）\n\n'
        '3. 语音识别功能：\n'
        '   - 使用 Web Speech API (webkitSpeechRecognition)\n'
        '   - "按住说话"按钮：mousedown 开始录音识别，mouseup 停止\n'
        '   - 识别结果显示在输入框中，同时自动发送\n\n'
        '4. WebSocket 连接：\n'
        '   - 连接 ws://localhost:8000/ws\n'
        '   - 发送消息格式：{type:"chat", text:"用户文字", image:"base64图片(可选)"}\n'
        '   - 接收消息格式：{type:"reply", text:"AI回复文字"}\n'
        '   - 接收到AI回复后显示在对话区，并用 SpeechSynthesis 朗读出来\n\n'
        '5. 自动截帧：\n'
        '   - 每次用户发送消息时，自动截取当前摄像头画面一并发送\n'
        '   - 如果摄像头未开启，则只发文字'
    )

    add_prompt_block(doc,
        '帮我优化前端的语音交互体验：\n'
        '1. "按住说话"按钮按住时显示录音动画（红色脉冲圆点）\n'
        '2. 识别到的文字实时显示在按钮上方（ interim results）\n'
        '3. AI回复时显示"思考中..."的加载动画\n'
        '4. 消息气泡区分用户（右侧蓝色）和AI（左侧灰色）\n'
        '5. AI回复朗读时在消息旁显示🔊动画'
    )

    doc.add_page_break()

    # ====================================================================
    # STEP 3: Backend WebSocket + DeepSeek
    # ====================================================================
    add_heading_styled(doc, '第三步：后端 WebSocket 与 DeepSeek API 接入', 1)

    p = doc.add_paragraph()
    run = p.add_run('⏱ 预计耗时：1.5-2小时')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE0, 0x4B, 0x00)

    add_heading_styled(doc, '3.1 要做的事情', 2)
    tasks = [
        '实现 WebSocket 消息接收与路由',
        '接入 DeepSeek Chat API（文字对话）',
        '接入 DeepSeek Vision API（图片理解）',
        '实现多轮对话上下文管理',
        '将 AI 回复通过 WebSocket 返回前端',
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    add_heading_styled(doc, '3.2 DeepSeek API 关键信息', 2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    make_header_row(table, ['项目', '说明'])
    ds_rows = [
        ['API Base URL', 'https://api.deepseek.com'],
        ['Chat 模型', 'deepseek-chat（文字对话）'],
        ['Vision 模型', 'deepseek-reasoner 或 deepseek-chat（支持图片）'],
        ['兼容格式', 'OpenAI Chat Completions API 格式'],
        ['SDK', 'pip install openai（用 openai SDK 调用）'],
        ['图片传入', 'messages 中 content 为数组，包含 {type:"image_url", image_url:{url:"data:image/jpeg;base64,..."}}'],
        ['API Key 位置', 'https://platform.deepseek.com/api_keys'],
        ['免费额度', '新用户有赠送额度，足够开发测试'],
    ]
    for r in ds_rows:
        add_row(table, r)

    add_heading_styled(doc, '3.3 问 Claude 的语句', 2)

    add_prompt_block(doc,
        '帮我编写 backend/ai_service.py，封装 DeepSeek API 调用：\n\n'
        '1. 使用 openai SDK 调用 DeepSeek API（兼容 OpenAI 格式）\n'
        '   - base_url = "https://api.deepseek.com"\n'
        '   - api_key 从环境变量 DEEPSEEK_API_KEY 读取\n\n'
        '2. 实现 ChatService 类：\n'
        '   - __init__(): 初始化 OpenAI client\n'
        '   - chat(text: str, history: list) -> str: 纯文字对话\n'
        '   - chat_with_image(text: str, image_base64: str, history: list) -> str: '
        '图文对话，图片以 base64 格式传入\n'
        '   - 内部维护 messages 列表实现多轮对话上下文\n'
        '   - 系统提示词设置为："你是一个AI视觉助手，用户会通过摄像头向你展示事物，'
        '你需要仔细观察图片内容并回答用户的问题。回答要简洁、准确、友好。如果图片中有文字请帮忙识别。"\n\n'
        '3. 对话历史管理：\n'
        '   - 每个 WebSocket 连接维护独立的对话历史\n'
        '   - 最多保留最近 20 条消息（10轮对话）\n'
        '   - 超出时裁剪最早的消息（保留 system prompt）'
    )

    add_prompt_block(doc,
        '帮我修改 backend/main.py，实现完整的 WebSocket 对话服务：\n\n'
        '1. WebSocket 端点 /ws：\n'
        '   - 接收 JSON 消息：{type:"chat", text:"用户说的话", image:"base64图片(可选)"}\n'
        '   - 如果有 image 字段，调用 chat_with_image()\n'
        '   - 如果只有 text 字段，调用 chat()\n'
        '   - 返回：{type:"reply", text:"AI回复"}\n\n'
        '2. 每个 WebSocket 连接创建独立的 ChatService 实例\n\n'
        '3. 添加错误处理：\n'
        '   - API 调用超时（30秒）返回友好提示\n'
        '   - API 调用失败返回错误信息\n'
        '   - WebSocket 断开时清理资源\n\n'
        '4. 添加日志：打印收到的消息类型、AI响应时间\n\n'
        '5. CORS 配置允许前端访问'
    )

    add_prompt_block(doc,
        '帮我测试 DeepSeek API 调用是否正常：\n'
        '在 backend/ 下创建 test_api.py，\n'
        '1. 测试纯文字对话：发送"你好，请介绍一下自己"\n'
        '2. 测试图片理解：找一张测试图片的 base64，发送"请描述这张图片"\n'
        '3. 打印请求耗时和返回结果\n'
        '4. 如果失败，打印详细的错误信息'
    )

    doc.add_page_break()

    # ====================================================================
    # STEP 4: Full Integration
    # ====================================================================
    add_heading_styled(doc, '第四步：前后端联调 — 跑通完整链路', 1)

    p = doc.add_paragraph()
    run = p.add_run('⏱ 预计耗时：1-1.5小时')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE0, 0x4B, 0x00)

    add_heading_styled(doc, '4.1 要做的事情', 2)
    tasks = [
        '启动后端服务，确认 WebSocket 端点可用',
        '前端连接 WebSocket，发送文字消息，验证AI回复',
        '前端截帧+文字一起发送，验证视觉理解功能',
        '语音识别 → 发送 → AI回复 → TTS朗读，完整闭环测试',
        '处理各种异常情况（断连、超时、无摄像头等）',
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    add_heading_styled(doc, '4.2 联调检查清单', 2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    make_header_row(table, ['测试项', '操作', '预期结果'])
    check_rows = [
        ['健康检查', '浏览器访问 http://localhost:8000/', '返回 {"status":"ok"}'],
        ['WebSocket连接', '打开前端页面', '控制台显示"已连接"'],
        ['纯文字对话', '输入"你好"并发送', 'AI返回问候语'],
        ['图片+文字', '开启摄像头，发送"你看到什么"', 'AI描述摄像头画面'],
        ['语音输入', '按住说话按钮说"你好"', '识别文字→发送→收到AI回复→语音朗读'],
        ['多轮对话', '连续问几个问题', 'AI能记住上下文'],
        ['断线重连', '断开网络再恢复', '自动重连成功'],
    ]
    for r in check_rows:
        add_row(table, r)

    add_heading_styled(doc, '4.3 问 Claude 的语句', 2)

    add_prompt_block(doc,
        '现在前后端代码都有了，请帮我做联调检查：\n\n'
        '1. 检查 backend/main.py 的 WebSocket 路径是否和前端 app.js 中的连接地址一致\n'
        '2. 检查消息格式是否匹配（前端发送的 JSON 字段名和后端解析的是否一致）\n'
        '3. 检查图片 base64 传输是否正确（是否包含 data:image/jpeg;base64, 前缀）\n'
        '4. 添加 WebSocket 连接状态显示（前端页面顶部显示连接状态绿/红灯）\n'
        '5. 添加自动重连机制（断开后每3秒重试）\n'
        '6. 如果有任何不一致的地方，请帮我修复'
    )

    add_prompt_block(doc,
        '帮我处理前端的异常情况：\n'
        '1. 摄像头权限被拒绝时，显示友好提示"请允许摄像头权限"，不阻塞其他功能\n'
        '2. WebSocket 连接失败时，显示"服务未启动，请先启动后端"\n'
        '3. AI回复超时（15秒）时，显示"AI思考超时，请重试"\n'
        '4. 麦克风不可用时，隐藏语音按钮，只保留文字输入\n'
        '5. 添加一个"重新连接"按钮'
    )

    doc.add_page_break()

    # ====================================================================
    # STEP 5: Cost Optimization
    # ====================================================================
    add_heading_styled(doc, '第五步：成本优化 — 帧采样与 VAD', 1)

    p = doc.add_paragraph()
    run = p.add_run('⏱ 预计耗时：1-1.5小时')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE0, 0x4B, 0x00)

    add_heading_styled(doc, '5.1 要做的事情', 2)
    tasks = [
        '实现视频帧智能采样（画面不变不重发）',
        '实现前端 VAD（静音时不发送音频）',
        '实现响应缓存（相似画面复用结果）',
        '添加费用统计显示',
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    add_heading_styled(doc, '5.2 成本控制核心逻辑', 2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    make_header_row(table, ['策略', '实现位置', '原理'])
    cost_rows = [
        ['帧差异检测', '前端 Canvas', '比较相邻帧像素，变化<阈值时不截帧'],
        ['定时采样兜底', '前端 setInterval', '最长5秒必须采样一次'],
        ['VAD静音检测', '前端 Web Audio API', '音频能量<阈值时不发送'],
        ['对话历史裁剪', '后端 Python', '只保留最近10轮对话'],
        ['图片压缩', '前端 Canvas', 'JPEG质量0.6，长边不超过720px'],
        ['相同图片跳过', '前端哈希比较', '连续相同帧不重复发送'],
    ]
    for r in cost_rows:
        add_row(table, r)

    add_heading_styled(doc, '5.3 问 Claude 的语句', 2)

    add_prompt_block(doc,
        '帮我实现前端的视频帧智能采样功能，修改 app.js：\n\n'
        '1. 实现帧差异检测：\n'
        '   - 维护一个 previousFrame Canvas 存储上一次发送的帧\n'
        '   - captureFrame() 时，先计算当前帧与上一帧的像素差异（MSE算法）\n'
        '   - 差异 < 阈值(比如15) 时，返回 null 表示不需要重新发送\n'
        '   - 差异 >= 阈值时，更新 previousFrame 并返回 base64\n\n'
        '2. 自动采样：\n'
        '   - 用户开启"自动模式"后，每2秒自动截帧检测\n'
        '   - 如果画面有变化，自动发送给AI（附带"请描述当前画面的变化"）\n'
        '   - 如果无变化，不发送\n\n'
        '3. 在页面上显示一个开关："自动观察模式"\n'
        '   - 开启后AI持续观察画面\n'
        '   - 显示当前采样状态（采样中/静止/检测到变化）'
    )

    add_prompt_block(doc,
        '帮我实现一个简单的费用统计功能：\n\n'
        '后端 (main.py)：\n'
        '1. 每次调用 DeepSeek API 后，记录使用的 token 数\n'
        '2. 添加 GET /api/stats 接口，返回：{total_tokens, total_calls, estimated_cost}\n'
        '3. DeepSeek 价格：输入 1元/百万token，输出 2元/百万token（大致估算）\n\n'
        '前端 (app.js)：\n'
        '1. 页面底部显示一个小的状态栏：API调用次数 | Token用量 | 预估费用\n'
        '2. 每次收到AI回复后更新统计\n'
        '3. 费用超过阈值时变红提醒'
    )

    doc.add_page_break()

    # ====================================================================
    # STEP 6: UI Polish + Advanced Features
    # ====================================================================
    add_heading_styled(doc, '第六步：UI美化与功能增强', 1)

    p = doc.add_paragraph()
    run = p.add_run('⏱ 预计耗时：1.5-2小时')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE0, 0x4B, 0x00)

    add_heading_styled(doc, '6.1 要做的事情', 2)
    tasks = [
        'UI 整体美化（深色主题、动画、布局优化）',
        '添加 OCR 文字识别模式',
        '添加"翻译"快捷指令',
        '添加截图保存功能',
        '添加对话历史导出',
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    add_heading_styled(doc, '6.2 快捷指令设计', 2)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    make_header_row(table, ['指令', '触发方式', '效果'])
    cmd_rows = [
        ['识别文字', '点击"OCR"按钮或说"帮我识别文字"', '发送当前帧 + "请识别图中的所有文字"'],
        ['翻译', '点击"翻译"按钮或说"帮我翻译"', '发送当前帧 + "请识别图中的文字并翻译成中文"'],
        ['描述场景', '点击"描述"按钮或说"这是哪里"', '发送当前帧 + "请详细描述这个场景"'],
        ['识别物体', '点击"识别"按钮或说"这是什么"', '发送当前帧 + "请识别图中的物体并介绍"'],
        ['截图保存', '点击"截图"按钮', '保存当前帧+AI回复为图片'],
    ]
    for r in cmd_rows:
        add_row(table, r)

    add_heading_styled(doc, '6.3 问 Claude 的语句', 2)

    add_prompt_block(doc,
        '帮我美化前端界面，要求：\n\n'
        '1. 深色主题：背景 #1a1a2e，卡片 #16213e，强调色 #0f3460，高亮 #e94560\n'
        '2. 消息气泡样式：\n'
        '   - 用户消息：右对齐，蓝色渐变背景(#0f3460→#1a56db)，白色文字\n'
        '   - AI消息：左对齐，深灰背景(#2d2d44)，白色文字\n'
        '   - 圆角气泡，带头像图标（用户👤 机器人🤖）\n'
        '3. 视频区：圆角边框，角落显示"LIVE"红色指示灯\n'
        '4. 按钮：圆角，hover 时有光晕效果\n'
        '5. AI"思考中"状态：显示三个跳动的点动画\n'
        '6. 整体使用 CSS Grid 或 Flexbox 布局，现代感强\n'
        '7. 所有动画使用 CSS transition/animation，不用JS动画库'
    )

    add_prompt_block(doc,
        '帮我添加快捷指令按钮功能：\n\n'
        '在对话输入框上方添加一排快捷按钮：\n'
        '[📷 OCR识别] [🌐 翻译] [🔍 识别物体] [📝 描述场景] [💾 截图]\n\n'
        '点击每个按钮时：\n'
        '1. 自动截取当前摄像头画面\n'
        '2. 发送预设的指令文字 + 图片\n'
        '3. OCR识别："请识别图片中的所有文字，保持原始排版"\n'
        '4. 翻译："请识别图片中的文字，如果是外文请翻译成中文"\n'
        '5. 识别物体："请识别图片中的主要物体，给出名称和简要介绍"\n'
        '6. 描述场景："请详细描述这个场景，包括环境、人物、物品等"\n'
        '7. 截图：将当前帧+最近一条AI回复保存为图片下载\n\n'
        '按钮样式：半透明胶囊按钮，横向排列，可滚动'
    )

    doc.add_page_break()

    # ====================================================================
    # STEP 7: Testing + Docs
    # ====================================================================
    add_heading_styled(doc, '第七步：测试、文档与演示准备', 1)

    p = doc.add_paragraph()
    run = p.add_run('⏱ 预计耗时：1-1.5小时')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE0, 0x4B, 0x00)

    add_heading_styled(doc, '7.1 要做的事情', 2)
    tasks = [
        '完整功能测试',
        '录制演示视频（3-5分钟）',
        '更新设计文档（填入实际实现情况）',
        '撰写项目总结',
    ]
    for t in tasks:
        doc.add_paragraph(t, style='List Bullet')

    add_heading_styled(doc, '7.2 演示脚本', 2)

    demo_steps = [
        '开场（30秒）：介绍项目名称、技术栈、核心功能',
        '基础对话（1分钟）：展示文字对话、语音对话',
        '视觉理解（1分钟）：展示物体识别、场景描述、OCR文字识别',
        '成本控制（1分钟）：展示智能采样、费用统计面板',
        '高级功能（30秒）：翻译、截图、快捷指令',
        '总结（30秒）：成本控制策略总结、未来改进方向',
    ]
    for i, step in enumerate(demo_steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_heading_styled(doc, '7.3 问 Claude 的语句', 2)

    add_prompt_block(doc,
        '帮我做项目的最终检查：\n\n'
        '1. 检查所有文件是否有语法错误\n'
        '2. 检查 WebSocket 消息格式前后端是否一致\n'
        '3. 检查 DeepSeek API 调用参数是否正确（model名称、图片格式等）\n'
        '4. 检查 .env 文件是否有正确的 API Key\n'
        '5. 列出所有需要手动测试的功能点\n'
        '6. 如果发现问题，直接帮我修复'
    )

    add_prompt_block(doc,
        '帮我更新设计文档 docs/02-设计文档-用户故事.docx 中的实现情况：\n\n'
        '根据我们实际开发的功能，帮我生成一段文字，说明：\n'
        '1. 计划实现的用户故事有哪些，最终实现了哪些\n'
        '2. 未实现的故事及原因\n'
        '3. 额外实现的功能（计划外的）\n\n'
        '同时帮我生成成本控制文档 docs/05-成本控制设计文档.docx 中的实测数据部分：\n'
        '1. 实际采用了哪些成本控制策略\n'
        '2. 每个策略的实现方式和效果\n'
        '3. 优化前后的API调用次数对比（估算值）'
    )

    doc.add_page_break()

    # ====================================================================
    # APPENDIX: Quick Reference
    # ====================================================================
    add_heading_styled(doc, '附录：开发速查表', 1)

    add_heading_styled(doc, 'A. 启动命令速查', 2)

    cmds = [
        ('后端启动', 'cd eye-talk/backend && pip install -r requirements.txt && python main.py'),
        ('前端启动', 'cd eye-talk/frontend && 用浏览器直接打开 index.html（或 python -m http.server 3000）'),
        ('查看后端日志', '终端中直接查看 uvicorn 输出'),
        ('测试API', 'python test_api.py'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    make_header_row(table, ['操作', '命令'])
    for r in cmds:
        add_row(table, r)

    add_heading_styled(doc, 'B. DeepSeek API 调用示例代码', 2)

    code = '''# 纯文字对话
from openai import OpenAI

client = OpenAI(
    api_key="your_deepseek_api_key",
    base_url="https://api.deepseek.com"
)

response = client.chat.completions.create(
    model="deepseek-chat",
    messages=[
        {"role": "system", "content": "你是一个AI视觉助手"},
        {"role": "user", "content": "你好"}
    ]
)
print(response.choices[0].message.content)

# 图文对话（视觉理解）
response = client.chat.completions.create(
    model="deepseek-chat",
    messages=[
        {"role": "system", "content": "你是一个AI视觉助手"},
        {"role": "user", "content": [
            {"type": "text", "text": "请描述这张图片"},
            {"type": "image_url", "image_url": {
                "url": "data:image/jpeg;base64,/9j/4AAQ..."
            }}
        ]}
    ]
)'''

    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

    add_heading_styled(doc, 'C. 常见问题排查', 2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    make_header_row(table, ['问题', '解决方案'])
    faq = [
        ['摄像头打不开', '检查浏览器权限设置，确保使用 HTTPS 或 localhost'],
        ['WebSocket 连不上', '确认后端已启动，检查端口和路径是否一致'],
        ['DeepSeek API 报错', '检查 API Key 是否正确，账户是否有余额'],
        ['语音识别不工作', 'Web Speech API 仅支持 Chrome/Edge，需 HTTPS 或 localhost'],
        ['图片太大传输慢', '降低 Canvas JPEG 质量到 0.5-0.6，缩小分辨率到 480p'],
        ['AI回复太慢', '减少图片分辨率、缩短对话历史、设置 max_tokens'],
        ['TTS没声音', '检查浏览器是否允许自动播放，需用户先交互一次'],
    ]
    for r in faq:
        add_row(table, r)

    # ===== SAVE =====
    path = os.path.join(BASE_DIR, 'docs', '06-一天半冲刺开发计划.docx')
    doc.save(path)
    print(f"Generated: {path}")


if __name__ == '__main__':
    generate()
